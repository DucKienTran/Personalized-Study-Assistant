import asyncio
from dataclasses import dataclass
import logging
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import fitz
import pymupdf4llm

from app.exceptions import AppError, BadRequestError, InternalServerError
from app.storage.base import StorageService

logger = logging.getLogger(__name__)
CHARS_PER_PAGE = 3000  # Ngưỡng ước lượng số ký tự trên một trang A4 fluid


@dataclass(slots=True)
class ParsedDocument:
    markdown: str
    total_pages: int


class DocumentParserService:
    async def parse_document(
        self,
        object_name: str,
        storage_service: StorageService,
    ) -> ParsedDocument:
        temp_path = await storage_service.download_temp_file(object_name)
        try:
            return await asyncio.to_thread(self._parse_sync, temp_path)
        except AppError:
            raise
        except Exception:
            logger.exception("[Parser] Failed to parse document: %s", object_name)
            raise InternalServerError("Failed to parse the uploaded document.")
        finally:
            await storage_service.delete_temp_file(temp_path)

    def _parse_sync(self, file_path: str) -> ParsedDocument:
        extension = Path(file_path).suffix.lower()
        if extension == ".pdf":
            return self._parse_pdf(file_path)
        if extension == ".docx":
            return self._parse_docx(file_path)
        raise BadRequestError("Only PDF and DOCX documents are currently supported.")

    def _parse_pdf(self, pdf_path: str) -> ParsedDocument:
        doc = fitz.open(pdf_path)
        try:
            total_pages = len(doc)
        finally:
            doc.close()

        pages = pymupdf4llm.to_markdown(pdf_path, page_chunks=True)

        if not pages:
            raise BadRequestError("Unable to extract text from the PDF document.")

        md_metadata = pages[0].get("metadata", {})
        if "page" not in md_metadata and "page_number" not in md_metadata:
            logger.warning(
                "[Parser] pymupdf4llm metadata keys: %s — page marker sẽ bị bỏ qua",
                list(md_metadata.keys()),
            )

        parts = []
        for p in pages:
            meta = p.get("metadata", {})
            # Thư viện trả về index từ 0, ta cộng 1 để ra số trang thực tế của con người (1, 2, 3...)
            raw_page = meta.get("page", meta.get("page_number"))
            page_num = raw_page if raw_page is not None else None
            
            marker = f"<!--page:{page_num}-->\n" if page_num is not None else ""
            parts.append(f"{marker}{p['text']}")

        markdown = "".join(parts)

        if not markdown.strip():
            raise BadRequestError("Unable to extract text from the PDF document.")

        return ParsedDocument(markdown=markdown, total_pages=total_pages)

    def _parse_docx(self, docx_path: str) -> ParsedDocument:
        document = Document(docx_path)
        markdown_lines: list[str] = []
        
        current_page = 1
        accumulated_chars = 0

        # Bơm marker trang 1 vào ngay đầu file DOCX
        markdown_lines.append(f"<!--page:{current_page}-->")

        # Duyệt qua các phần tử để đảm bảo giữ nguyên thứ tự xuất hiện trong file
        for child in document.element.body:
            element_text = ""
            
            if child.tag.endswith("p"):
                paragraph = Paragraph(child, document)
                text = paragraph.text.strip()
                if not text:
                    continue

                style = paragraph.style.name.lower().strip() if paragraph.style else ""
                if style.startswith("heading 1"):
                    element_text = f"# {text}\n"
                elif style.startswith("heading 2"):
                    element_text = f"## {text}\n"
                elif style.startswith("heading 3"):
                    element_text = f"### {text}\n"
                elif style.startswith("heading 4"):
                    element_text = f"#### {text}\n"
                else:
                    element_text = f"{text}\n"

            elif child.tag.endswith("tbl"):
                table = Table(child, document)
                table_lines = []
                for i, row in enumerate(table.rows):
                    row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    table_lines.append("| " + " | ".join(row_text) + " |")
                    if i == 0:
                        table_lines.append("| " + " | ".join(["---"] * len(row.cells)) + " |")
                element_text = "\n".join(table_lines) + "\n"

            if element_text:
                markdown_lines.append(element_text)
                accumulated_chars += len(element_text)
                
                if accumulated_chars >= CHARS_PER_PAGE:
                    current_page += 1
                    markdown_lines.append(f"<!--page:{current_page}-->")
                    accumulated_chars = 0

        markdown = "\n".join(markdown_lines)
        if not markdown.strip():
            raise BadRequestError("Unable to extract text from the DOCX document.")

        return ParsedDocument(markdown=markdown, total_pages=current_page)